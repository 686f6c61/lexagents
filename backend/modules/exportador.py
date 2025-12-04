# -*- coding: utf-8 -*-
"""
LexAgents - Sistema Multi-Agente de Extracción Legal
https://github.com/686f6c61/lexagents

Módulo Exportador de Resultados
Exporta referencias y resultados a diferentes formatos:
- Markdown (.md)
- Texto plano (.txt)
- Word (.docx)

Author: 686f6c61
Version: 0.2.0
License: MIT
"""

import logging
import re
from typing import Dict, List, Any, Optional
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from modules.boe_article_fetcher import get_boe_article_fetcher
from modules.legal_abbreviations import (
    expandir_sigla,
    es_sigla_conocida,
    es_referencia_contextual
)
from modules.eurlex_fetcher import enriquecer_referencia_eurlex
from agents.eurlex_article_extractor_agent import EurlexArticleExtractorAgent

logger = logging.getLogger(__name__)


class Exportador:
    """
    Exportador de resultados a múltiples formatos

    Genera documentos profesionales con:
    - Referencias extraídas
    - Títulos completos de leyes
    - Textos completos de artículos del BOE
    - Enlaces al BOE
    """

    def __init__(self, output_dir: str = None):
        """
        Inicializa el exportador

        Args:
            output_dir: Directorio de salida (default: ../data/results)
        """
        if output_dir is None:
            base_path = Path(__file__).parent.parent.parent
            output_dir = base_path / "data" / "results"

        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Inicializar BOE article fetcher
        self.boe_fetcher = get_boe_article_fetcher()

        # Inicializar EUR-Lex article extractor
        self.eurlex_extractor = EurlexArticleExtractorAgent()

        logger.info(f"✅ Exportador inicializado - Output: {self.output_dir}")

    @staticmethod
    def _enriquecer_eurlex(texto_completo: str) -> Optional[Dict]:
        """
        Enriquece una referencia EUR-Lex con datos completos

        Args:
            texto_completo: Texto de la referencia (ej: "Reglamento (UE) 2017/1939")

        Returns:
            Dict con datos EUR-Lex (CELEX, URLs, título) o None
        """
        try:
            return enriquecer_referencia_eurlex(texto_completo)
        except Exception as e:
            logger.error(f"Error enriqueciendo EUR-Lex: {e}")
            return None

    @staticmethod
    def _debe_exportarse(ref: Dict) -> bool:
        """
        Determina si una referencia debe exportarse

        Regla: SOLO exportar referencias que tienen BOE-ID (validado por el agente)
        o son legislación europea (EUR-Lex).

        El BOE-ID lo infiere el agente validador buscando en la API del BOE,
        no tiene que estar explícito en el texto original.

        Args:
            ref: Referencia a evaluar

        Returns:
            True si debe exportarse
        """
        # Si tiene BOE-ID (inferido por el agente validador) → SÍ exportar
        if ref.get('boe_id'):
            return True

        # Si es legislación europea → SÍ exportar (se enriquecerá con EUR-Lex)
        texto_completo = ref.get('texto_completo', '')
        es_eurlex = any(palabra in texto_completo for palabra in [
            'Reglamento (UE)', 'Reglamento (CE)', 'Reglamento UE', 'Reglamento CE',
            'Directiva (UE)', 'Directiva (CE)', 'Directiva UE', 'Directiva CE',
            'Decisión (UE)', 'Decisión (CE)'
        ])

        if es_eurlex:
            return True

        # Sin BOE-ID ni EUR-Lex → NO exportar
        return False

    @staticmethod
    def _expandir_nombre_ley(nombre_ley: str) -> str:
        """
        Expande siglas de leyes a nombres completos

        Args:
            nombre_ley: Nombre de la ley (puede ser sigla)

        Returns:
            Nombre expandido
        """
        if es_sigla_conocida(nombre_ley):
            return expandir_sigla(nombre_ley)
        return nombre_ley

    def exportar_todo(
        self,
        referencias: List[Dict],
        informe_auditoria: Dict = None,
        metricas: Dict = None,
        tema: str = "tema",
        formatos: List[str] = None,
        referencias_inferidas: List[Dict] = None
    ) -> Dict[str, str]:
        """
        Exporta a todos los formatos solicitados

        Args:
            referencias: Referencias verificadas a exportar
            informe_auditoria: Informe de auditoría (IGNORADO - no se usa)
            metricas: Métricas del pipeline (IGNORADO - no se usa)
            tema: Nombre del tema (usado como nombre de archivo)
            formatos: Lista de formatos ['md', 'txt', 'docx', 'pdf'] (default: md, txt, docx)
            referencias_inferidas: Referencias inferidas (BETA) - opcional

        Returns:
            Dict con rutas de archivos generados
        """
        if formatos is None:
            formatos = ['md', 'txt', 'docx']  # PDF es opcional

        if referencias_inferidas is None:
            referencias_inferidas = []

        logger.info(f"📤 Exportando resultados a {len(formatos)} formatos")
        if referencias_inferidas:
            logger.info(f"   + {len(referencias_inferidas)} referencias inferidas (BETA)")

        # Enriquecer referencias con títulos de leyes y textos de artículos
        referencias_enriquecidas = self._enriquecer_referencias(referencias)

        # También enriquecer referencias inferidas
        referencias_inferidas_enriquecidas = self._enriquecer_referencias(referencias_inferidas) if referencias_inferidas else []

        archivos_generados = {}

        # Nombre base de archivo (usar directamente el nombre del JSON sin timestamp)
        nombre_base = self._sanitizar_nombre(tema)

        # Exportar cada formato
        if 'md' in formatos:
            archivo_md = self.exportar_markdown(
                referencias_enriquecidas, nombre_base, tema,
                referencias_inferidas=referencias_inferidas_enriquecidas
            )
            archivos_generados['markdown'] = archivo_md

        if 'txt' in formatos:
            archivo_txt = self.exportar_texto(
                referencias_enriquecidas, nombre_base, tema,
                referencias_inferidas=referencias_inferidas_enriquecidas
            )
            archivos_generados['texto'] = archivo_txt

        if 'docx' in formatos:
            archivo_docx = self.exportar_word(
                referencias_enriquecidas, nombre_base, tema,
                referencias_inferidas=referencias_inferidas_enriquecidas
            )
            archivos_generados['word'] = archivo_docx

        if 'pdf' in formatos:
            archivo_pdf = self.exportar_pdf(
                referencias_enriquecidas, nombre_base, tema,
                referencias_inferidas=referencias_inferidas_enriquecidas
            )
            archivos_generados['pdf'] = archivo_pdf

        logger.info(f"✅ Exportación completada: {len(archivos_generados)} archivos")

        return archivos_generados

    def _enriquecer_referencias(self, referencias: List[Dict]) -> List[Dict]:
        """
        Enriquece las referencias con títulos de leyes y textos de artículos del BOE
        También filtra referencias que no deben exportarse

        Args:
            referencias: Lista de referencias

        Returns:
            Lista de referencias enriquecidas y filtradas
        """
        logger.info(f"🔍 Enriqueciendo {len(referencias)} referencias con datos del BOE...")

        referencias_enriquecidas = []

        for ref in referencias:
            # Filtrar: solo exportar referencias específicas
            if not self._debe_exportarse(ref):
                logger.debug(f"   Saltando referencia completa: {ref.get('texto_completo', 'N/A')}")
                continue

            ref_enriquecida = ref.copy()

            # Expandir siglas en nombre de ley
            if ref.get('ley_nombre'):
                nombre_expandido = self._expandir_nombre_ley(ref['ley_nombre'])
                if nombre_expandido != ref['ley_nombre']:
                    ref_enriquecida['_ley_nombre_expandido'] = nombre_expandido
                    logger.debug(f"   Expandido: {ref['ley_nombre']} → {nombre_expandido}")

            # Detectar y enriquecer legislación europea (Reglamentos, Directivas, Decisiones UE/CE)
            texto_completo = ref.get('texto_completo', '')
            es_eurlex = any(palabra in texto_completo for palabra in [
                'Reglamento (UE)', 'Reglamento (CE)', 'Reglamento UE', 'Reglamento CE',
                'Directiva (UE)', 'Directiva (CE)', 'Directiva UE', 'Directiva CE',
                'Decisión (UE)', 'Decisión (CE)'
            ])

            if es_eurlex:
                logger.debug(f"   Detectado acto EUR-Lex: {texto_completo[:80]}...")
                datos_eurlex = self._enriquecer_eurlex(texto_completo)

                if datos_eurlex:
                    # Añadir datos EUR-Lex a la referencia
                    celex = datos_eurlex.get('celex')
                    ref_enriquecida['_eurlex_celex'] = celex
                    ref_enriquecida['_eurlex_url'] = datos_eurlex.get('urls', {}).get('principal')

                    # Título completo del EUR-Lex (si está disponible)
                    if datos_eurlex.get('titulo_completo'):
                        ref_enriquecida['_eurlex_titulo'] = datos_eurlex['titulo_completo']

                    logger.debug(f"   ✅ EUR-Lex enriquecido: {celex}")

                    # Si es un artículo específico, extraer su texto completo
                    if ref.get('articulo') and celex:
                        try:
                            articulo_str = str(ref['articulo'])

                            # Detectar múltiples artículos (ej: "22 y 25", "22, 25")
                            articulos_a_extraer = []
                            if ' y ' in articulo_str or ',' in articulo_str:
                                # Múltiples artículos
                                import re
                                articulos_a_extraer = re.findall(r'\d+', articulo_str)
                                logger.debug(f"   🔍 Detectados múltiples artículos EUR-Lex: {articulos_a_extraer}")
                            else:
                                # Un solo artículo
                                articulos_a_extraer = [articulo_str]

                            # Extraer cada artículo
                            textos_articulos = []
                            titulos_articulos = []

                            for num_art in articulos_a_extraer[:3]:  # Límite de 3 artículos para no saturar
                                logger.debug(f"   🔍 Extrayendo artículo {num_art} de {celex}...")
                                resultado_articulo = self.eurlex_extractor.procesar({
                                    'celex': celex,
                                    'articulo': num_art,
                                    'idioma': 'ES'
                                })

                                if resultado_articulo.get('exito'):
                                    textos_articulos.append(resultado_articulo.get('texto_completo', ''))
                                    titulos_articulos.append(f"Art. {num_art}: {resultado_articulo.get('titulo_articulo', '')}")
                                    logger.debug(f"   ✅ Artículo EUR-Lex {num_art} extraído ({len(resultado_articulo.get('texto_completo', ''))} chars)")
                                else:
                                    logger.debug(f"   ⚠️  No se pudo extraer artículo EUR-Lex {num_art}: {resultado_articulo.get('error')}")

                            # Guardar todos los textos
                            if textos_articulos:
                                ref_enriquecida['_eurlex_texto_articulo'] = '\n\n---\n\n'.join(textos_articulos)
                                ref_enriquecida['_eurlex_titulo_articulo'] = ' | '.join(titulos_articulos)
                                logger.info(f"   ✅ {len(textos_articulos)} artículo(s) EUR-Lex extraído(s)")

                        except Exception as e:
                            logger.error(f"   ❌ Error extrayendo artículo EUR-Lex: {e}")
                else:
                    logger.debug(f"   ⚠️  No se pudo enriquecer EUR-Lex")

            # Obtener título completo de la ley si tiene boe_id
            if ref.get('boe_id'):
                titulo_ley = self.boe_fetcher.obtener_titulo_ley(ref['boe_id'])
                if titulo_ley:
                    ref_enriquecida['_titulo_ley'] = titulo_ley

                # Si es un artículo específico, obtener su texto completo
                if ref.get('articulo'):
                    logger.info(f"   📥 Obteniendo texto de artículo {ref['articulo']} de {ref['boe_id']}")
                    articulo_data = self.boe_fetcher.obtener_articulo(
                        ref['boe_id'],
                        ref['articulo']
                    )
                    if articulo_data:
                        ref_enriquecida['_texto_articulo'] = articulo_data.get('texto', '')
                        ref_enriquecida['_titulo_articulo'] = articulo_data.get('titulo', '')
                        if articulo_data.get('es_subapartado'):
                            logger.info(f"   ✅ Texto obtenido (artículo base para subapartado {articulo_data.get('numero_subapartado')})")
                        else:
                            logger.info(f"   ✅ Texto obtenido del artículo {ref['articulo']}")
                    else:
                        logger.warning(f"   ❌ NO se pudo obtener texto del artículo {ref['articulo']} de {ref['boe_id']}")
                        logger.warning(f"      Referencia: {ref.get('texto_completo', 'N/A')}")

                # Si tiene lista de artículos (referencias inferidas), obtener texto del primero
                elif ref.get('articulos') and isinstance(ref['articulos'], list) and len(ref['articulos']) > 0:
                    primer_articulo = ref['articulos'][0]
                    logger.info(f"   📥 Obteniendo texto del primer artículo inferido: {primer_articulo} de {ref['boe_id']}")
                    articulo_data = self.boe_fetcher.obtener_articulo(
                        ref['boe_id'],
                        primer_articulo
                    )
                    if articulo_data:
                        ref_enriquecida['_texto_articulo_ejemplo'] = articulo_data.get('texto', '')
                        ref_enriquecida['_titulo_articulo_ejemplo'] = articulo_data.get('titulo', '')
                        ref_enriquecida['_numero_articulo_ejemplo'] = primer_articulo
                        logger.info(f"   ✅ Texto del primer artículo obtenido (ejemplo para referencia inferida)")
                    else:
                        logger.warning(f"   ❌ NO se pudo obtener texto del primer artículo inferido {primer_articulo}")

            referencias_enriquecidas.append(ref_enriquecida)

        logger.info(f"✅ {len(referencias_enriquecidas)} de {len(referencias)} referencias filtradas y enriquecidas")
        return referencias_enriquecidas

    def exportar_markdown(
        self,
        referencias: List[Dict],
        nombre_archivo: str = "referencias",
        tema: str = "Tema",
        referencias_inferidas: List[Dict] = None
    ) -> str:
        """
        Exporta a Markdown

        Args:
            referencias: Referencias verificadas
            nombre_archivo: Nombre base del archivo
            tema: Nombre del tema
            referencias_inferidas: Referencias inferidas (BETA) - opcional

        Returns:
            Ruta del archivo generado
        """
        filepath = self.output_dir / f"{nombre_archivo}.md"

        contenido = []

        # Título
        contenido.append("# REFERENCIAS LEGALES EXTRAÍDAS")
        contenido.append("")
        contenido.append(f"**Tema:** {tema}")
        contenido.append(f"**Fecha:** {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        contenido.append(f"**Total referencias verificadas:** {len(referencias)}")
        if referencias_inferidas:
            contenido.append(f"**Total referencias inferidas (BETA):** {len(referencias_inferidas)}")
        contenido.append("")
        contenido.append("---")
        contenido.append("")

        # SECCIÓN 1: Referencias Verificadas
        contenido.append("# 📋 SECCIÓN 1: REFERENCIAS VERIFICADAS")
        contenido.append("")
        contenido.append("Estas referencias han sido extraídas directamente del texto y validadas contra el BOE oficial.")
        contenido.append("")
        contenido.append("---")
        contenido.append("")

        # Referencias verificadas
        for i, ref in enumerate(referencias, 1):
            contenido.append(f"## {i}. {ref.get('texto_completo', 'N/A')}")
            contenido.append("")

            # Título completo de la ley
            if ref.get('_titulo_ley'):
                contenido.append(f"**{ref['_titulo_ley']}**")
                contenido.append("")
            elif ref.get('_ley_nombre_expandido'):
                contenido.append(f"**{ref['_ley_nombre_expandido']}**")
                contenido.append("")

            # Tipo de referencia
            tipo = ref.get('tipo', '').lower()
            if 'capitulo' in tipo or 'capítulo' in tipo:
                contenido.append("**Tipo:** Capítulo")
                contenido.append("")
            elif 'titulo' in tipo or 'título' in tipo:
                contenido.append("**Tipo:** Título de código")
                contenido.append("")

            # BOE ID y URL
            if ref.get('boe_id'):
                contenido.append(f"**BOE-ID:** {ref['boe_id']}")
                boe_url = f"https://www.boe.es/buscar/act.php?id={ref['boe_id']}"
                contenido.append(f"**BOE URL:** [{boe_url}]({boe_url})")
                contenido.append("")

            # EUR-Lex
            if ref.get('_eurlex_url'):
                if ref.get('_eurlex_celex'):
                    contenido.append(f"**CELEX:** {ref['_eurlex_celex']}")
                contenido.append(f"**EUR-Lex:** [{ref['_eurlex_url']}]({ref['_eurlex_url']})")
                if ref.get('_eurlex_titulo'):
                    contenido.append(f"**Título:** {ref['_eurlex_titulo']}")
                contenido.append("")

            # Texto del artículo
            texto_articulo = ref.get('_eurlex_texto_articulo') or ref.get('_texto_articulo')
            titulo_articulo = ref.get('_eurlex_titulo_articulo') or ref.get('_titulo_articulo')

            if texto_articulo:
                contenido.append("### Texto del Artículo")
                contenido.append("")
                if titulo_articulo:
                    contenido.append(f"**{titulo_articulo}**")
                    contenido.append("")
                texto_limpio = self._limpiar_html(texto_articulo)
                contenido.append(texto_limpio)
                contenido.append("")
            contenido.append("---")
            contenido.append("")

        # SECCIÓN 2: Referencias Inferidas (BETA)
        if referencias_inferidas:
            contenido.append("")
            contenido.append("# 🧠 SECCIÓN 2: POSIBLES NORMAS INFERIDAS (BETA)")
            contenido.append("")
            contenido.append("⚠️ **IMPORTANTE:** Estas referencias han sido sugeridas por IA basándose en conceptos legales detectados en el texto.")
            contenido.append("")
            contenido.append("**No fueron mencionadas explícitamente** en el material de estudio, pero pueden ser relevantes para el tema.")
            contenido.append("")
            contenido.append("✅ **Recomendación:** Verifica estas referencias con tu material de estudio antes de incluirlas.")
            contenido.append("")
            contenido.append("---")
            contenido.append("")

            for i, ref in enumerate(referencias_inferidas, 1):
                # Para referencias inferidas, construir el texto completo
                ley = ref.get('ley', 'Ley desconocida')
                articulos = ref.get('articulos', [])
                concepto = ref.get('concepto_detectado', '')
                confianza = ref.get('confianza', 0)

                contenido.append(f"## BETA-{i}. {ley}")
                contenido.append("")
                contenido.append(f"**Concepto detectado:** {concepto}")
                contenido.append(f"**Confianza IA:** {confianza}%")
                contenido.append(f"**Artículos sugeridos:** {', '.join(articulos[:20])}{'...' if len(articulos) > 20 else ''}")
                contenido.append("")

                # BOE ID y URL
                if ref.get('boe_id'):
                    contenido.append(f"**BOE-ID:** {ref['boe_id']}")
                    boe_url = f"https://www.boe.es/buscar/act.php?id={ref['boe_id']}"
                    contenido.append(f"**BOE URL:** [{boe_url}]({boe_url})")
                    contenido.append("")

                # Si se enriqueció con textos de artículos
                if ref.get('_texto_articulo'):
                    contenido.append("### Texto del Primer Artículo (Ejemplo)")
                    contenido.append("")
                    texto_limpio = self._limpiar_html(ref['_texto_articulo'])
                    # Limitar longitud
                    if len(texto_limpio) > 500:
                        texto_limpio = texto_limpio[:500] + "..."
                    contenido.append(texto_limpio)
                    contenido.append("")

                contenido.append("---")
                contenido.append("")

        # Escribir archivo
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(contenido))

        logger.info(f"✅ Markdown generado: {filepath.name}")
        return str(filepath)

    def exportar_texto(
        self,
        referencias: List[Dict],
        nombre_archivo: str = "referencias",
        tema: str = "Tema",
        referencias_inferidas: List[Dict] = None
    ) -> str:
        """
        Exporta a texto plano

        Args:
            referencias: Referencias verificadas
            nombre_archivo: Nombre base del archivo
            tema: Nombre del tema
            referencias_inferidas: Referencias inferidas (BETA) - opcional

        Returns:
            Ruta del archivo generado
        """
        filepath = self.output_dir / f"{nombre_archivo}.txt"

        contenido = []

        # Título
        contenido.append("=" * 80)
        contenido.append("REFERENCIAS LEGALES EXTRAÍDAS")
        contenido.append("=" * 80)
        contenido.append("")
        contenido.append(f"Tema: {tema}")
        contenido.append(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        contenido.append(f"Total referencias verificadas: {len(referencias)}")
        if referencias_inferidas:
            contenido.append(f"Total referencias inferidas (BETA): {len(referencias_inferidas)}")
        contenido.append("")
        contenido.append("=" * 80)
        contenido.append("")

        # SECCIÓN 1: Referencias Verificadas
        contenido.append("SECCIÓN 1: REFERENCIAS VERIFICADAS")
        contenido.append("")
        contenido.append("Estas referencias han sido extraídas directamente del texto y validadas")
        contenido.append("contra el BOE oficial.")
        contenido.append("")
        contenido.append("=" * 80)
        contenido.append("")

        # Referencias verificadas
        for i, ref in enumerate(referencias, 1):
            contenido.append("-" * 80)
            contenido.append(f"{i}. {ref.get('texto_completo', 'N/A')}")
            contenido.append("-" * 80)
            contenido.append("")

            # Título completo de la ley
            if ref.get('_titulo_ley'):
                contenido.append(f"LEY: {ref['_titulo_ley']}")
                contenido.append("")
            elif ref.get('_ley_nombre_expandido'):
                contenido.append(f"LEY: {ref['_ley_nombre_expandido']}")
                contenido.append("")

            # Tipo de referencia
            tipo = ref.get('tipo', '').lower()
            if 'capitulo' in tipo or 'capítulo' in tipo:
                contenido.append("TIPO: Capítulo")
                contenido.append("")
            elif 'titulo' in tipo or 'título' in tipo:
                contenido.append("TIPO: Título de código")
                contenido.append("")

            # BOE ID y URL
            if ref.get('boe_id'):
                contenido.append(f"BOE-ID: {ref['boe_id']}")
                contenido.append(f"BOE URL: https://www.boe.es/buscar/act.php?id={ref['boe_id']}")
                contenido.append("")

            # EUR-Lex
            if ref.get('_eurlex_url'):
                if ref.get('_eurlex_celex'):
                    contenido.append(f"CELEX: {ref['_eurlex_celex']}")
                contenido.append(f"EUR-Lex URL: {ref['_eurlex_url']}")
                if ref.get('_eurlex_titulo'):
                    contenido.append(f"Título EUR-Lex: {ref['_eurlex_titulo']}")
                contenido.append("")

            # Texto del artículo
            if ref.get('articulo'):
                texto_articulo = ref.get('_eurlex_texto_articulo') or ref.get('_texto_articulo')
                titulo_articulo = ref.get('_eurlex_titulo_articulo') or ref.get('_titulo_articulo')

                if texto_articulo:
                    contenido.append("TEXTO DEL ARTÍCULO:")
                    contenido.append("")
                    if titulo_articulo:
                        contenido.append(titulo_articulo)
                        contenido.append("")
                    texto_limpio = self._limpiar_html(texto_articulo)
                    contenido.append(texto_limpio)
                    contenido.append("")

            contenido.append("")

        # SECCIÓN 2: Referencias Inferidas (BETA)
        if referencias_inferidas:
            contenido.append("")
            contenido.append("=" * 80)
            contenido.append("SECCIÓN 2: POSIBLES NORMAS INFERIDAS (BETA)")
            contenido.append("=" * 80)
            contenido.append("")
            contenido.append("IMPORTANTE: Estas referencias han sido sugeridas por IA basándose en")
            contenido.append("conceptos legales detectados en el texto.")
            contenido.append("")
            contenido.append("No fueron mencionadas explícitamente en el material de estudio, pero")
            contenido.append("pueden ser relevantes para el tema.")
            contenido.append("")
            contenido.append("RECOMENDACIÓN: Verifica estas referencias con tu material de estudio")
            contenido.append("antes de incluirlas.")
            contenido.append("")
            contenido.append("=" * 80)
            contenido.append("")

            for i, ref in enumerate(referencias_inferidas, 1):
                ley = ref.get('ley', 'Ley desconocida')
                articulos = ref.get('articulos', [])
                concepto = ref.get('concepto_detectado', '')
                confianza = ref.get('confianza', 0)

                contenido.append("-" * 80)
                contenido.append(f"BETA-{i}. {ley}")
                contenido.append("-" * 80)
                contenido.append("")
                contenido.append(f"Concepto detectado: {concepto}")
                contenido.append(f"Confianza IA: {confianza}%")
                contenido.append(f"Artículos sugeridos: {', '.join(articulos[:20])}{'...' if len(articulos) > 20 else ''}")
                contenido.append("")

                # BOE ID y URL
                if ref.get('boe_id'):
                    contenido.append(f"BOE-ID: {ref['boe_id']}")
                    contenido.append(f"BOE URL: https://www.boe.es/buscar/act.php?id={ref['boe_id']}")
                    contenido.append("")

                contenido.append("")

        contenido.append("=" * 80)

        # Escribir archivo
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(contenido))

        logger.info(f"✅ Texto generado: {filepath.name}")
        return str(filepath)

    def exportar_word(
        self,
        referencias: List[Dict],
        nombre_archivo: str = "referencias",
        tema: str = "Tema",
        referencias_inferidas: List[Dict] = None
    ) -> str:
        """
        Exporta a Word (DOCX)

        Args:
            referencias: Referencias verificadas
            nombre_archivo: Nombre base del archivo
            tema: Nombre del tema
            referencias_inferidas: Referencias inferidas (BETA) - opcional

        Returns:
            Ruta del archivo generado
        """
        filepath = self.output_dir / f"{nombre_archivo}.docx"

        doc = Document()

        # Configurar estilo del documento
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0, 0, 0)

        # Título principal
        titulo = doc.add_heading('REFERENCIAS LEGALES EXTRAÍDAS', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titulo.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Metadatos
        p = doc.add_paragraph()
        run = p.add_run(f"Tema: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p.add_run(tema)
        run.font.color.rgb = RGBColor(0, 0, 0)

        p = doc.add_paragraph()
        run = p.add_run(f"Fecha: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p.add_run(datetime.now().strftime('%d/%m/%Y %H:%M'))
        run.font.color.rgb = RGBColor(0, 0, 0)

        p = doc.add_paragraph()
        run = p.add_run(f"Total referencias verificadas: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p.add_run(str(len(referencias)))
        run.font.color.rgb = RGBColor(0, 0, 0)

        if referencias_inferidas:
            p = doc.add_paragraph()
            run = p.add_run(f"Total referencias inferidas (BETA): ")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            run = p.add_run(str(len(referencias_inferidas)))
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Línea separadora
        doc.add_paragraph("_" * 80)

        # SECCIÓN 1: Referencias Verificadas
        seccion1 = doc.add_heading('SECCIÓN 1: REFERENCIAS VERIFICADAS', 1)
        for run in seccion1.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

        p = doc.add_paragraph()
        run = p.add_run('Estas referencias han sido extraídas directamente del texto y validadas contra el BOE oficial.')
        run.font.color.rgb = RGBColor(0, 0, 0)

        doc.add_paragraph("_" * 80)

        # Referencias
        for i, ref in enumerate(referencias, 1):
            # Número y texto de referencia
            heading = doc.add_heading(f"{i}. {ref.get('texto_completo', 'N/A')}", 2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Título completo de la ley
            if ref.get('_titulo_ley'):
                p = doc.add_paragraph()
                run = p.add_run(ref['_titulo_ley'])
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif ref.get('_ley_nombre_expandido'):
                p = doc.add_paragraph()
                run = p.add_run(ref['_ley_nombre_expandido'])
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Tipo de referencia
            tipo = ref.get('tipo', '').lower()
            if 'capitulo' in tipo or 'capítulo' in tipo:
                p = doc.add_paragraph()
                run = p.add_run('Tipo: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run = p.add_run('Capítulo')
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif 'titulo' in tipo or 'título' in tipo:
                p = doc.add_paragraph()
                run = p.add_run('Tipo: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run = p.add_run('Título de código')
                run.font.color.rgb = RGBColor(0, 0, 0)

            # BOE ID y URL
            if ref.get('boe_id'):
                p = doc.add_paragraph()
                run = p.add_run('BOE-ID: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run = p.add_run(ref['boe_id'])
                run.font.color.rgb = RGBColor(0, 0, 0)

                p = doc.add_paragraph()
                run = p.add_run('BOE URL: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                boe_url = f"https://www.boe.es/buscar/act.php?id={ref['boe_id']}"
                run = p.add_run(boe_url)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.underline = True

            # EUR-Lex (Reglamentos/Directivas/Decisiones UE)
            if ref.get('_eurlex_url'):
                # CELEX
                if ref.get('_eurlex_celex'):
                    p = doc.add_paragraph()
                    run = p.add_run('CELEX: ')
                    run.bold = True
                    run = p.add_run(ref['_eurlex_celex'])

                # URL EUR-Lex
                p = doc.add_paragraph()
                run = p.add_run('EUR-Lex URL: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run = p.add_run(ref['_eurlex_url'])
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.underline = True

                # Título EUR-Lex (si está disponible)
                if ref.get('_eurlex_titulo'):
                    p = doc.add_paragraph()
                    run = p.add_run('Título EUR-Lex: ')
                    run.bold = True
                    run = p.add_run(ref['_eurlex_titulo'])

            # Texto del artículo (si es artículo específico)
            if ref.get('articulo'):  # Es un artículo específico
                subheading = doc.add_heading('Texto del Artículo', 3)
                for run in subheading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)

                texto_articulo = ref.get('_eurlex_texto_articulo') or ref.get('_texto_articulo')
                titulo_articulo = ref.get('_eurlex_titulo_articulo') or ref.get('_titulo_articulo')

                if texto_articulo:
                    # Texto disponible (EUR-Lex o BOE)
                    if titulo_articulo:
                        p_titulo = doc.add_paragraph()
                        run = p_titulo.add_run(titulo_articulo)
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)

                    texto_limpio = self._limpiar_html(texto_articulo)
                    p = doc.add_paragraph(texto_limpio)
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    # Texto no disponible - dirigir al BOE/EUR-Lex
                    p = doc.add_paragraph()
                    run = p.add_run('El texto completo de este artículo no está disponible en formato estructurado en la API del BOE.')
                    run.font.color.rgb = RGBColor(0, 0, 0)

                    p = doc.add_paragraph()
                    run = p.add_run('Puede consultarlo directamente en:')
                    run.font.color.rgb = RGBColor(0, 0, 0)

                    if ref.get('boe_id'):
                        p = doc.add_paragraph()
                        boe_url = f"https://www.boe.es/buscar/act.php?id={ref['boe_id']}"
                        run = p.add_run(boe_url)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.underline = True

            # Espacio entre referencias
            doc.add_paragraph()

        # SECCIÓN 2: Referencias Inferidas (BETA)
        if referencias_inferidas:
            doc.add_page_break()

            # Título sección BETA
            seccion2 = doc.add_heading('SECCIÓN 2: POSIBLES NORMAS INFERIDAS (BETA)', 1)
            for run in seccion2.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Advertencia
            p = doc.add_paragraph()
            run = p.add_run('⚠️ IMPORTANTE: ')
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            run = p.add_run('Estas referencias han sido sugeridas por IA basándose en conceptos legales detectados en el texto.')
            run.font.color.rgb = RGBColor(0, 0, 0)

            p = doc.add_paragraph()
            run = p.add_run('No fueron mencionadas explícitamente')
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            run = p.add_run(' en el material de estudio, pero pueden ser relevantes para el tema.')
            run.font.color.rgb = RGBColor(0, 0, 0)

            p = doc.add_paragraph()
            run = p.add_run('✅ RECOMENDACIÓN: ')
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            run = p.add_run('Verifica estas referencias con tu material de estudio antes de incluirlas.')
            run.font.color.rgb = RGBColor(0, 0, 0)

            doc.add_paragraph("_" * 80)

            # Referencias inferidas
            for i, ref in enumerate(referencias_inferidas, 1):
                ley = ref.get('ley', 'Ley desconocida')
                articulos = ref.get('articulos', [])
                concepto = ref.get('concepto_detectado', '')
                confianza = ref.get('confianza', 0)

                # Número y ley
                heading = doc.add_heading(f"BETA-{i}. {ley}", 2)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)

                # Concepto detectado
                p = doc.add_paragraph()
                run = p.add_run('Concepto detectado: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run = p.add_run(concepto)
                run.font.color.rgb = RGBColor(0, 0, 0)

                # Confianza
                p = doc.add_paragraph()
                run = p.add_run('Confianza IA: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run = p.add_run(f"{confianza}%")
                run.font.color.rgb = RGBColor(0, 0, 0)

                # Artículos sugeridos
                p = doc.add_paragraph()
                run = p.add_run('Artículos sugeridos: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                articulos_texto = ', '.join(articulos[:20])
                if len(articulos) > 20:
                    articulos_texto += '...'
                run = p.add_run(articulos_texto)
                run.font.color.rgb = RGBColor(0, 0, 0)

                # BOE ID y URL
                if ref.get('boe_id'):
                    p = doc.add_paragraph()
                    run = p.add_run('BOE-ID: ')
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run = p.add_run(ref['boe_id'])
                    run.font.color.rgb = RGBColor(0, 0, 0)

                    p = doc.add_paragraph()
                    run = p.add_run('BOE URL: ')
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    boe_url = f"https://www.boe.es/buscar/act.php?id={ref['boe_id']}"
                    run = p.add_run(boe_url)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.underline = True

                # Texto del artículo ejemplo (si está disponible)
                if ref.get('_texto_articulo_ejemplo'):
                    subheading = doc.add_heading('Ejemplo de artículo', 3)
                    for run in subheading.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)

                    if ref.get('_titulo_articulo_ejemplo'):
                        p_titulo = doc.add_paragraph()
                        run = p_titulo.add_run(f"Art. {ref.get('_numero_articulo_ejemplo')}: {ref['_titulo_articulo_ejemplo']}")
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)

                    texto_limpio = self._limpiar_html(ref['_texto_articulo_ejemplo'])
                    p = doc.add_paragraph(texto_limpio)
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)

                # Espacio entre referencias
                doc.add_paragraph()

        # Guardar documento
        doc.save(filepath)

        logger.info(f"✅ Word generado: {filepath.name}")
        return str(filepath)

    def _limpiar_html(self, html: str) -> str:
        """
        Limpia HTML básico para convertirlo a texto plano

        Args:
            html: Contenido HTML

        Returns:
            Texto limpio
        """
        import re
        from html import unescape

        # Remover etiquetas HTML
        texto = re.sub(r'<[^>]+>', '', html)
        # Decodificar entidades HTML
        texto = unescape(texto)
        # Limpiar espacios múltiples
        texto = re.sub(r'\s+', ' ', texto)
        # Limpiar líneas vacías múltiples
        texto = re.sub(r'\n\s*\n', '\n\n', texto)

        return texto.strip()

    def exportar_pdf(
        self,
        referencias: List[Dict],
        nombre_archivo: str = "referencias",
        tema: str = "Tema",
        referencias_inferidas: List[Dict] = None
    ) -> str:
        """
        Exporta a PDF usando reportlab

        Args:
            referencias: Referencias verificadas
            nombre_archivo: Nombre base del archivo
            tema: Nombre del tema
            referencias_inferidas: Referencias inferidas (BETA) - opcional

        Returns:
            Ruta del archivo generado
        """
        filepath = self.output_dir / f"{nombre_archivo}.pdf"

        # Crear documento PDF
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )

        # Estilos
        styles = getSampleStyleSheet()

        # Estilo para título
        title_style = ParagraphStyle(
            'TituloCustom',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=12,
            alignment=TA_CENTER
        )

        # Estilo para subtítulos
        subtitle_style = ParagraphStyle(
            'SubtituloCustom',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=8
        )

        # Estilo para texto normal
        normal_style = ParagraphStyle(
            'NormalCustom',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6
        )

        # Estilo para sección BETA
        beta_style = ParagraphStyle(
            'BETACustom',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#d97706'),
            spaceAfter=8
        )

        # Contenido del documento
        story = []

        # Título
        story.append(Paragraph(f"<b>Referencias Legales: {tema}</b>", title_style))
        story.append(Spacer(1, 0.5*cm))

        # Información del documento
        fecha = datetime.now().strftime("%d/%m/%Y %H:%M")
        story.append(Paragraph(f"<i>Generado: {fecha}</i>", normal_style))
        story.append(Paragraph(f"<i>Total de referencias verificadas: {len(referencias)}</i>", normal_style))
        if referencias_inferidas:
            story.append(Paragraph(f"<i>Total de referencias inferidas (BETA): {len(referencias_inferidas)}</i>", normal_style))
        story.append(Spacer(1, 0.5*cm))

        # === REFERENCIAS VERIFICADAS ===
        story.append(Paragraph("<b>Referencias Verificadas</b>", subtitle_style))
        story.append(Spacer(1, 0.3*cm))

        for idx, ref in enumerate(referencias, 1):
            # Texto completo de la referencia (como en Word)
            texto_ref = ref.get('texto_completo', 'N/A')
            story.append(Paragraph(f"<b>{idx}. {texto_ref}</b>", subtitle_style))

            # Título completo de la ley
            if ref.get('_titulo_ley'):
                story.append(Paragraph(f"<b>{ref['_titulo_ley']}</b>", normal_style))
            elif ref.get('_ley_nombre_expandido'):
                story.append(Paragraph(f"<b>{ref['_ley_nombre_expandido']}</b>", normal_style))

            # Tipo de referencia
            tipo = ref.get('tipo', '').lower()
            if 'capitulo' in tipo or 'capítulo' in tipo:
                story.append(Paragraph(f"<b>Tipo:</b> Capítulo", normal_style))
            elif 'titulo' in tipo or 'título' in tipo:
                story.append(Paragraph(f"<b>Tipo:</b> Título de código", normal_style))

            # BOE ID y URL
            if ref.get('boe_id'):
                story.append(Paragraph(f"<b>BOE-ID:</b> {ref['boe_id']}", normal_style))
                boe_url = f"https://www.boe.es/buscar/act.php?id={ref['boe_id']}"
                story.append(Paragraph(f"<b>BOE URL:</b> <link href='{boe_url}'>{boe_url}</link>", normal_style))

            # EUR-Lex (como en Word)
            if ref.get('_eurlex_url'):
                if ref.get('_eurlex_celex'):
                    story.append(Paragraph(f"<b>CELEX:</b> {ref['_eurlex_celex']}", normal_style))
                eurlex_url = ref['_eurlex_url']
                story.append(Paragraph(f"<b>EUR-Lex URL:</b> <link href='{eurlex_url}'>{eurlex_url}</link>", normal_style))
                if ref.get('_eurlex_titulo'):
                    story.append(Paragraph(f"<b>Título EUR-Lex:</b> {ref['_eurlex_titulo']}", normal_style))

            # Texto del artículo (si es artículo específico) - Como en Word
            if ref.get('articulo'):
                story.append(Paragraph("<b>Texto del Artículo</b>", subtitle_style))

                texto_articulo = ref.get('_eurlex_texto_articulo') or ref.get('_texto_articulo')
                titulo_articulo = ref.get('_eurlex_titulo_articulo') or ref.get('_titulo_articulo')

                if texto_articulo:
                    if titulo_articulo:
                        story.append(Paragraph(f"<b>{titulo_articulo}</b>", normal_style))

                    texto_limpio = self._limpiar_html(texto_articulo)
                    story.append(Paragraph(texto_limpio, normal_style))
                else:
                    story.append(Paragraph("El texto completo de este artículo no está disponible en formato estructurado en la API del BOE.", normal_style))
                    story.append(Paragraph("Puede consultarlo directamente en:", normal_style))
                    if ref.get('boe_id'):
                        boe_url = f"https://www.boe.es/buscar/act.php?id={ref['boe_id']}"
                        story.append(Paragraph(f"<link href='{boe_url}'>{boe_url}</link>", normal_style))

            story.append(Spacer(1, 0.4*cm))

        # === REFERENCIAS INFERIDAS (BETA) ===
        if referencias_inferidas:
            story.append(PageBreak())
            story.append(Paragraph("<b>Posibles Normas Inferidas (BETA)</b>", beta_style))
            story.append(Spacer(1, 0.3*cm))

            # Advertencia
            advertencia = (
                "<b>IMPORTANTE:</b> Estas referencias han sido sugeridas por IA basándose "
                "en conceptos legales mencionados en el documento. Requieren verificación "
                "manual antes de ser utilizadas."
            )
            story.append(Paragraph(advertencia, normal_style))
            story.append(Spacer(1, 0.4*cm))

            for idx, ref in enumerate(referencias_inferidas, 1):
                # Concepto detectado
                concepto = ref.get('concepto_detectado', 'Desconocido')
                story.append(Paragraph(f"<b>{idx}. Concepto: {concepto}</b>", normal_style))

                # Ley sugerida
                ley = ref.get('ley', 'Ley desconocida')
                story.append(Paragraph(f"<b>Ley sugerida:</b> {ley}", normal_style))

                # BOE ID
                if ref.get('boe_id'):
                    story.append(Paragraph(f"<b>BOE:</b> {ref['boe_id']}", normal_style))

                # Artículos
                articulos = ref.get('articulos', [])
                if articulos:
                    arts_texto = ', '.join([str(a) for a in articulos[:10]])
                    if len(articulos) > 10:
                        arts_texto += f' ... (+{len(articulos)-10} más)'
                    story.append(Paragraph(f"<b>Artículos:</b> {arts_texto}", normal_style))

                # Confianza
                confianza = ref.get('confianza', 0)
                story.append(Paragraph(f"<b>Confianza IA:</b> {confianza}%", normal_style))

                # Texto del artículo ejemplo (si está disponible)
                if ref.get('_texto_articulo_ejemplo'):
                    story.append(Spacer(1, 0.2*cm))
                    story.append(Paragraph("<b>Ejemplo de artículo:</b>", subtitle_style))
                    if ref.get('_titulo_articulo_ejemplo'):
                        story.append(Paragraph(
                            f"<b>Art. {ref.get('_numero_articulo_ejemplo')}:</b> {ref['_titulo_articulo_ejemplo']}",
                            normal_style
                        ))
                    texto_limpio = self._limpiar_html(ref['_texto_articulo_ejemplo'])
                    story.append(Paragraph(texto_limpio, normal_style))

                story.append(Spacer(1, 0.4*cm))

        # Generar PDF
        doc.build(story)

        logger.info(f"✅ PDF generado: {filepath}")

        return str(filepath)

    def _sanitizar_nombre(self, nombre: str) -> str:
        """
        Sanitiza un nombre de archivo

        Args:
            nombre: Nombre original

        Returns:
            Nombre sanitizado
        """
        # Quitar caracteres no permitidos
        nombre = nombre.replace(' ', '_')
        nombre = ''.join(c for c in nombre if c.isalnum() or c in ['_', '-'])
        return nombre[:50]  # Limitar longitud


# Función helper
def exportar_resultados(
    referencias: List[Dict],
    informe_auditoria: Dict = None,
    metricas: Dict = None,
    tema: str = "tema",
    formatos: List[str] = None,
    output_dir: str = None
) -> Dict[str, str]:
    """
    Función helper para exportar resultados

    Args:
        referencias: Referencias a exportar
        informe_auditoria: Informe de auditoría (IGNORADO)
        metricas: Métricas (IGNORADO)
        tema: Nombre del tema
        formatos: Formatos a exportar
        output_dir: Directorio de salida

    Returns:
        Dict con rutas de archivos generados
    """
    exportador = Exportador(output_dir)
    return exportador.exportar_todo(
        referencias,
        informe_auditoria,
        metricas,
        tema,
        formatos
    )
