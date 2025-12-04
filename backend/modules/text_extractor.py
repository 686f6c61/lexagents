# -*- coding: utf-8 -*-
"""
LexAgents - Sistema Multi-Agente de Extracción Legal
https://github.com/686f6c61/lexagents

Módulo de Extracción de Texto Multi-Formato
Extrae texto de diferentes formatos de archivo:
- PDF (.pdf) usando PyMuPDF
- Word (.docx) usando python-docx
- Texto plano (.txt)

Author: 686f6c61
Version: 0.2.0
License: MIT
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional
import json

logger = logging.getLogger(__name__)


class TextExtractor:
    """
    Extractor de texto multi-formato

    Soporta PDF, Word, TXT, MD y JSON
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.json'}

    def __init__(self):
        """Inicializa el extractor"""
        logger.info("✅ TextExtractor inicializado")

    def extraer_texto(self, file_path: str) -> str:
        """
        Extrae texto de un archivo según su extensión

        Args:
            file_path: Ruta al archivo

        Returns:
            Texto extraído del archivo

        Raises:
            ValueError: Si el formato no está soportado
            Exception: Si hay error en la extracción
        """
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Formato no soportado: {extension}. "
                f"Formatos válidos: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        logger.info(f"📄 Extrayendo texto de {path.name} ({extension})")

        try:
            if extension == '.pdf':
                return self._extraer_de_pdf(file_path)
            elif extension == '.docx':
                return self._extraer_de_word(file_path)
            elif extension == '.txt':
                return self._extraer_de_txt(file_path)
            elif extension == '.md':
                return self._extraer_de_markdown(file_path)
            elif extension == '.json':
                # JSON ya está en el formato correcto, no necesita extracción
                return None

        except Exception as e:
            logger.error(f"❌ Error extrayendo texto de {path.name}: {e}")
            raise

    def _extraer_de_pdf(self, file_path: str) -> str:
        """
        Extrae texto de un PDF usando PyPDF2

        Args:
            file_path: Ruta al archivo PDF

        Returns:
            Texto extraído
        """
        try:
            from PyPDF2 import PdfReader

            texto_completo = []

            # Abrir PDF
            reader = PdfReader(file_path)
            num_paginas = len(reader.pages)

            logger.info(f"   📖 PDF con {num_paginas} páginas")

            # Extraer texto de cada página
            for num_pagina, pagina in enumerate(reader.pages, 1):
                texto_pagina = pagina.extract_text()

                if texto_pagina and texto_pagina.strip():
                    texto_completo.append(texto_pagina)
                    logger.debug(f"   ✓ Página {num_pagina}: {len(texto_pagina)} caracteres")

            texto = "\n\n".join(texto_completo)
            logger.info(f"   ✅ Extraídos {len(texto)} caracteres del PDF")

            return texto

        except ImportError:
            raise ImportError(
                "PyPDF2 no está instalado. Ejecuta: pip install PyPDF2"
            )
        except Exception as e:
            logger.error(f"   ❌ Error extrayendo PDF: {e}")
            raise

    def _extraer_de_word(self, file_path: str) -> str:
        """
        Extrae texto de un documento Word (.docx)

        Args:
            file_path: Ruta al archivo Word

        Returns:
            Texto extraído
        """
        try:
            from docx import Document

            # Abrir documento
            doc = Document(file_path)

            # Extraer texto de todos los párrafos
            parrafos = []
            for parrafo in doc.paragraphs:
                texto = parrafo.text.strip()
                if texto:
                    parrafos.append(texto)

            # Extraer texto de tablas
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        texto = celda.text.strip()
                        if texto:
                            parrafos.append(texto)

            texto = "\n\n".join(parrafos)
            logger.info(f"   ✅ Extraídos {len(texto)} caracteres del Word")

            return texto

        except ImportError:
            raise ImportError(
                "python-docx no está instalado. Ejecuta: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"   ❌ Error extrayendo Word: {e}")
            raise

    def _extraer_de_txt(self, file_path: str) -> str:
        """
        Extrae texto de un archivo de texto plano

        Args:
            file_path: Ruta al archivo TXT

        Returns:
            Contenido del archivo
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                texto = f.read()

            logger.info(f"   ✅ Leídos {len(texto)} caracteres del TXT")
            return texto

        except UnicodeDecodeError:
            # Intentar con latin-1 si UTF-8 falla
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    texto = f.read()

                logger.warning("   ⚠️  Archivo codificado en latin-1, no UTF-8")
                logger.info(f"   ✅ Leídos {len(texto)} caracteres del TXT")
                return texto
            except Exception as e:
                logger.error(f"   ❌ Error leyendo TXT: {e}")
                raise

    def _extraer_de_markdown(self, file_path: str) -> str:
        """
        Extrae texto de un archivo Markdown

        Args:
            file_path: Ruta al archivo MD

        Returns:
            Contenido del archivo (se mantiene el markdown)
        """
        # El markdown se procesa igual que TXT
        # El pipeline puede manejar el formato markdown
        return self._extraer_de_txt(file_path)

    def convertir_a_json_tema(self, texto: str, nombre_archivo: str = "Documento") -> Dict[str, Any]:
        """
        Convierte texto plano al formato JSON esperado por el pipeline

        El formato JSON esperado es:
        {
            "document": {
                "documentVersion": {
                    "contenido": "<html>...</html>"
                }
            }
        }

        Args:
            texto: Texto extraído del documento
            nombre_archivo: Nombre del archivo original (para título)

        Returns:
            Diccionario con formato JSON del tema
        """
        # Escapar caracteres HTML si es necesario
        texto_escapado = texto.replace('&', '&amp;')
        texto_escapado = texto_escapado.replace('<', '&lt;')
        texto_escapado = texto_escapado.replace('>', '&gt;')

        # Crear HTML simple con párrafos
        # Dividir por líneas en blanco para párrafos
        lineas = texto.split('\n\n')
        parrafos_html = []

        for linea in lineas:
            linea_limpia = linea.strip()
            if linea_limpia:
                # Convertir saltos de línea simples en <br>
                linea_html = linea_limpia.replace('\n', '<br>')
                parrafos_html.append(f"<p>{linea_html}</p>")

        # Crear HTML completo
        html_contenido = f"""<div class="documento-estudiante">
<h1>{nombre_archivo}</h1>
{''.join(parrafos_html)}
</div>"""

        # Crear estructura JSON
        json_tema = {
            "document": {
                "documentVersion": {
                    "contenido": html_contenido
                }
            }
        }

        logger.info(f"   ✅ JSON creado: {len(html_contenido)} caracteres de contenido HTML")

        return json_tema

    def procesar_archivo(self, file_path: str, output_path: Optional[str] = None) -> str:
        """
        Procesa un archivo completo: extrae texto y genera JSON

        Args:
            file_path: Ruta al archivo de entrada
            output_path: Ruta opcional para guardar el JSON (si None, se genera automáticamente)

        Returns:
            Ruta del archivo JSON generado
        """
        path = Path(file_path)
        extension = path.suffix.lower()

        # Si ya es JSON, no necesita procesamiento
        if extension == '.json':
            logger.info(f"   ℹ️  Archivo ya es JSON, no requiere conversión")
            return file_path

        # Extraer texto
        texto = self.extraer_texto(file_path)

        if not texto or not texto.strip():
            raise ValueError(f"No se pudo extraer texto de {path.name}")

        # Convertir a JSON
        nombre_archivo = path.stem
        json_tema = self.convertir_a_json_tema(texto, nombre_archivo)

        # Determinar ruta de salida
        if output_path is None:
            output_path = path.parent / f"{nombre_archivo}.json"

        # Guardar JSON
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_tema, f, ensure_ascii=False, indent=2)

        logger.info(f"   ✅ JSON guardado en: {output_path}")

        return str(output_path)


# Función de conveniencia
def extraer_texto_de_archivo(file_path: str) -> str:
    """
    Función de conveniencia para extraer texto de un archivo

    Args:
        file_path: Ruta al archivo

    Returns:
        Texto extraído
    """
    extractor = TextExtractor()
    return extractor.extraer_texto(file_path)


def convertir_archivo_a_json(file_path: str, output_path: Optional[str] = None) -> str:
    """
    Función de conveniencia para convertir un archivo a JSON

    Args:
        file_path: Ruta al archivo de entrada
        output_path: Ruta opcional del JSON de salida

    Returns:
        Ruta del archivo JSON generado
    """
    extractor = TextExtractor()
    return extractor.procesar_archivo(file_path, output_path)
